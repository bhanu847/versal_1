import os
import uuid
import io
import logging
import pythoncom
import comtypes.client
import pdfplumber
import pytesseract
import cv2
import numpy as np
import fitz  # PyMuPDF
import pandas as pd
from flask import Flask, request, send_file, render_template_string, abort, redirect, url_for, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from PIL import Image
from docx import Document
import win32com.client as win32
from datetime import datetime


app = Flask(__name__)


def create_timestamped_upload_folder(base_folder):
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    folder_path = os.path.join(base_folder, timestamp)
    os.makedirs(folder_path, exist_ok=True)
    return folder_path




app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024


app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SAVED_FOLDER'] = 'saved'
app.config['THUMBNAIL_FOLDER'] = 'thumbnails'

for folder in app.config.values():
    if isinstance(folder, str):
        os.makedirs(folder, exist_ok=True)

logging.basicConfig(level=logging.INFO)



HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>PDF Tools Dashboard</title>
  <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
  <div class="container mx-auto px-4 py-8">
    <h1 class="text-4xl font-bold text-center text-blue-600 mb-10">RX PDF Tools</h1>
    <div class="grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 gap-6">

      {% for route, label, accept, name, extra_input in tools %}
      <div class="bg-white shadow rounded p-4">
        <h2 class="text-xl font-semibold text-blue-500 mb-2">{{ label }}</h2>

        {% if accept %}
          <form action="/{{ route }}" method="post" enctype="multipart/form-data">
            <input type="file" name="{{ name }}" accept="{{ accept }}" class="mb-2 w-full" {% if 'multiple' in extra_input %}multiple{% endif %} required>
            {{ extra_input | safe }}
            <button type="submit" class="w-full py-2 bg-blue-500 text-white rounded">Submit</button>
          </form>
        {% else %}
          <a href="/{{ route }}" class="w-full block text-center py-2 bg-purple-600 text-white rounded shadow hover:bg-purple-700">
             Open {{ label }}
          </a>
        {% endif %}
      </div>
      {% endfor %}

    </div>

    <div class="mt-8 text-center">
      <a href="/saved_files" class="text-blue-600 underline">📁 View Generate Files</a>
    </div>
    <div class="mt-2 text-center">
  <a href="/uploaded_files" class="text-green-600 underline">📤 View Uploaded Files</a>
</div>

  </div>
</body>
</html>
'''

SMART_SPLIT_TEMPLATE = """
<!DOCTYPE html>
<html lang=\"en\">
<head>
  <meta charset=\"UTF-8\">
  <title>Smart Split & Merge</title>
  <link rel=\"stylesheet\" href=\"https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css\">
  <style>
    body { background: #f7f7f7; }
    .panel { border: 1px solid #ddd; background: #fff; padding: 10px; height: 400px; overflow-y: auto; }
    .merge-box { border: 2px dashed #007bff; min-height: 100px; padding: 10px; }
    .pdf-preview { border: 1px solid #ccc; height: 400px; overflow-y: auto; background: #eee; }
    canvas { width: 100%; margin-bottom: 10px; }
  </style>
  <script src=\"https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.14.305/pdf.min.js\"></script>
</head>
<body class=\"container py-4\">
  <h2 class=\"text-center mb-4\">Smart PDF Split & Merge Tool</h2>
  <div class=\"mb-3\">
    <label for=\"pdfUpload\" class=\"form-label\">Upload PDF</label>
    <input type=\"file\" id=\"pdfUpload\" class=\"form-control\" accept=\"application/pdf\">
  </div>
  <div class=\"row\">
    <div class=\"col-md-3\">
      <h5>Available Pages</h5>
      <div class=\"panel\" id=\"availablePages\"></div>
    </div>
    <div class=\"col-md-6\">
      <h5>Selected Pages Preview</h5>
      <div class=\"pdf-preview\" id=\"selectedPreview\"></div>
    </div>
    <div class=\"col-md-3\">
      <h5>Pages to Merge</h5>
      <div class=\"merge-box\" id=\"mergeBox\"><p class=\"text-muted\">Selected pages will appear here</p></div>
    </div>
  </div>
  <hr>
  <div class=\"text-center my-3\">
    <button class=\"btn btn-primary\" id=\"generateBtn\">Generate PDFs</button>
  </div>
  <h5>Generated PDFs</h5>
  <ul class=\"list-group\" id=\"generatedList\"></ul>

  <div class=\"modal fade\" id=\"previewModal\" tabindex=\"-1\">
    <div class=\"modal-dialog modal-lg\">
      <div class=\"modal-content\">
        <div class=\"modal-header\"><h5 class=\"modal-title\">PDF Preview</h5><button type=\"button\" class=\"btn-close\" data-bs-dismiss=\"modal\"></button></div>
        <div class=\"modal-body\"><iframe id=\"modalPdfViewer\" src=\"\" width=\"100%\" height=\"500px\"></iframe></div>
      </div>
    </div>
  </div>

<script src=\"https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js\"></script>
<script>
let uploadedFilename = \"\";
let selectedPages = [];
let pdfUrl = \"\";

document.getElementById("pdfUpload").addEventListener("change", async function() {
  const file = this.files[0];
  if (!file) return;
  const formData = new FormData();
  formData.append("pdf", file);
  const res = await fetch("/smart_split_merge/upload", { method: "POST", body: formData });
  const data = await res.json();
  uploadedFilename = data.filename;
  pdfUrl = `/smart_split_merge/uploads/${uploadedFilename}`;
  const totalPages = data.total_pages;
  loadAvailablePages(totalPages);
  loadPdfThumbnails(pdfUrl, totalPages);
});

function loadAvailablePages(totalPages) {
  const container = document.getElementById("availablePages");
  container.innerHTML = "";
  for (let i = 1; i <= totalPages; i++) {
    const pageItem = document.createElement("div");
    pageItem.innerHTML = `<div class='form-check'><input class='form-check-input page-checkbox' type='checkbox' value='${i}' id='page${i}'><label class='form-check-label' for='page${i}'>Page ${i}</label></div>`;
    container.appendChild(pageItem);
  }
}

document.addEventListener("change", function(e) {
  if (e.target.classList.contains("page-checkbox")) {
    const mergeBox = document.getElementById("mergeBox");
    if (e.target.checked) {
      selectedPages.push(parseInt(e.target.value));
      let pageDiv = document.createElement("div");
      pageDiv.classList.add("badge", "bg-primary", "m-1");
      pageDiv.textContent = `Page ${e.target.value}`;
      pageDiv.id = "mergePage" + e.target.value;
      mergeBox.appendChild(pageDiv);
    } else {
      selectedPages = selectedPages.filter(p => p !== parseInt(e.target.value));
      let removeDiv = document.getElementById("mergePage" + e.target.value);
      if (removeDiv) removeDiv.remove();
    }
  }
});

async function loadPdfThumbnails(url, totalPages) {
  const pdf = await pdfjsLib.getDocument(url).promise;
  const previewContainer = document.getElementById("selectedPreview");
  previewContainer.innerHTML = "";
  for (let i = 1; i <= totalPages; i++) {
    const page = await pdf.getPage(i);
    const viewport = page.getViewport({ scale: 1.0 });
    const canvas = document.createElement("canvas");
    const ctx = canvas.getContext("2d");
    canvas.height = viewport.height;
    canvas.width = viewport.width;
    await page.render({ canvasContext: ctx, viewport: viewport }).promise;
    const wrapper = document.createElement("div");
    wrapper.classList.add("thumbnail-wrapper");
    wrapper.appendChild(canvas);
    previewContainer.appendChild(wrapper);
  }
}

document.getElementById("generateBtn").addEventListener("click", async function() {
  if (selectedPages.length === 0) { alert("Please select at least one page!"); return; }
  const res = await fetch("/smart_split_merge/generate", {
    method: "POST", headers: { "Content-Type": "application/json" },
    body: JSON.stringify({ filename: uploadedFilename, pages: selectedPages })
  });
  const data = await res.json();
  const generatedList = document.getElementById("generatedList");
  generatedList.innerHTML = "";
  data.generated.forEach(pdf => {
    let li = document.createElement("li");
    li.classList.add("list-group-item", "d-flex", "justify-content-between", "align-items-center");
    li.innerHTML = `<span>${pdf}</span><div><a class='btn btn-sm btn-success me-2' href='/smart_split_merge/generated/${pdf}' download>Download</a><button class='btn btn-sm btn-outline-primary preview-btn' data-pdf='/smart_split_merge/generated/${pdf}'>Preview</button></div>`;
    generatedList.appendChild(li);
  });
});

document.addEventListener("click", function(e) {
  if (e.target.classList.contains("preview-btn")) {
    const pdf = e.target.getAttribute("data-pdf");
    document.getElementById("modalPdfViewer").src = pdf;
    new bootstrap.Modal(document.getElementById("previewModal")).show();
  }
});
</script>
</body>
</html>
"""


def compress_pdf(input_path: str, output_path: str, dpi: int = 72, quality: int = 50):
    doc = fitz.open(input_path)
    images = []
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        img_byte = io.BytesIO()
        img.save(img_byte, format="JPEG", quality=quality)
        img_byte.seek(0)
        images.append(Image.open(img_byte).convert("RGB"))

    if not images:
        raise ValueError("No pages found in the PDF to compress.")

    images[0].save(output_path, save_all=True, append_images=images[1:])

import win32com.client as win32




# ==================== ROUTES ====================
@app.route('/')
def dashboard():
    tools = [
        ("compress", "Compress PDF", "application/pdf", "pdf", ""),
        ("merge", "Merge PDFs", "application/pdf", "pdfs", "<input type='file' name='pdfs' accept='application/pdf' multiple required>"),
        ("pdf_to_word", "PDF to Word", "application/pdf", "pdf", ""),
        ("pdf_to_excel", "PDF to Excel", "application/pdf", "pdf", ""),
        ("image_to_pdf", "Images to PDF", "image/*", "images", "<input type='file' name='images' accept='image/*' multiple required>"),
        ("ocr_image_to_docx", "OCR Image to Word", "image/*", "image", ""),
        ("remove_pages", "Remove Pages (PDF)", "application/pdf", "pdf", "<input name='pages' placeholder='e.g. 1,3,5' class='w-full mt-1 p-1 border rounded' required>"),
        ("rotate_pdf", "Rotate PDF Pages", "application/pdf", "pdf", "<input name='rotations' placeholder='e.g. 1:90,2:180' class='w-full mt-1 p-1 border rounded' required>"),

        ("smart_split_merge", "Smart Split & Merge", "", "", "")
    ]
    return render_template_string(HTML_TEMPLATE, tools=tools)

@app.route('/compress', methods=['POST'])
def compress():
    file = request.files['pdf']
    if not file:
        return "No PDF uploaded", 400
#####
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 100 * 1024 * 1024:
        return "File size should not exceed 100 MB", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])
    filename = secure_filename(file.filename)
    input_path = os.path.join(upload_folder, filename)
    file.save(input_path)



    output_filename = f"compressed_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)
    try:
        compress_pdf(input_path, output_path)
        return redirect(url_for('download_file', filename=output_filename))
    except Exception as e:
        return f"Compression failed: {e}", 500

@app.route('/merge', methods=['POST'])
def merge():
    files = request.files.getlist('pdfs')
    if len(files) < 2:
        return "At least two PDFs required.", 400

    total_size = sum([len(f.read()) for f in files])
    for f in files:
        f.seek(0)
    if total_size > 100 * 1024 * 1024:
        return "Combined file size should not exceed 100 MB", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])
    merger = PdfMerger()
    for file in files:
        filename = secure_filename(file.filename)
        input_path = os.path.join(upload_folder, filename)
        file.save(input_path)
        merger.append(input_path)

    output_filename = f"merged_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)
    merger.write(output_path)
    merger.close()
    return redirect(url_for('download_file', filename=output_filename))

#@app.route('/pdf_to_word', methods=['POST'])
#def pdf_to_word():
#   file = request.files['pdf']
#   if file:
#
#       file.seek(0, os.SEEK_END)
#        size = file.tell()
#        file.seek(0)
#        if size > 100 * 1024 * 1024:
#            return "File size should not exceed 100 MB", 400

#        upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

        # Save input
#        filename = secure_filename(file.filename)
#        input_path = os.path.join(upload_folder, filename)

#        file.save(input_path)
#        filename = f"converted_{uuid.uuid4().hex}.docx"

#        input_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
#        output_path = os.path.join(app.config['SAVED_FOLDER'], filename)
#        file.save(input_path)
#
#        pdf_document = fitz.open(input_path)
#        doc = Document()
#        for page_num in range(len(pdf_document)):
#            page = pdf_document.load_page(page_num)
#            text = page.get_text()
#            doc.add_paragraph(text)
#        doc.save(output_path)

#        return redirect(url_for('download_file', filename=filename))

@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
    file = request.files['pdf']
    if file:
        # Check file size (max 100MB)
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 100 * 1024 * 1024:
            return "File size should not exceed 100 MB", 400

        # Create timestamped upload folder
        upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

        # Save uploaded PDF
        filename = secure_filename(file.filename)
        input_path = os.path.join(upload_folder, filename)
        file.save(input_path)

        # Define output Word path
        output_filename = f"converted_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)

        # Convert PDF to Word
        pdf_document = fitz.open(input_path)
        doc = Document()
        for page in pdf_document:
            text = page.get_text()
            doc.add_paragraph(text)
        doc.save(output_path)

        return redirect(url_for('download_file', filename=output_filename))





@app.route('/pdf_to_excel', methods=['POST'])
def pdf_to_excel():
    file = request.files['pdf']
    if file:
#####
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 100 * 1024 * 1024:
            return "File size should not exceed 100 MB", 400

        upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

        # Save input
        filename = secure_filename(file.filename)
        input_path = os.path.join(upload_folder, filename)

        file.save(input_path)

####

        filename = f"excel_{uuid.uuid4().hex}.xlsx"
        output_path = os.path.join(app.config['SAVED_FOLDER'], filename)

        all_tables = []
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        df = pd.DataFrame(table)
                        all_tables.append(df)

        if all_tables:
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                for i, df in enumerate(all_tables):
                    df.to_excel(writer, sheet_name=f"Page_{i + 1}", index=False, header=False)
        else:
            df = pd.DataFrame([["No tables found in the PDF."]])
            df.to_excel(output_path, index=False, header=False)

        return redirect(url_for('download_file', filename=filename))

@app.route('/image_to_pdf', methods=['POST'])
def image_to_pdf():
    files = request.files.getlist('images')
    images = []
    for file in files:
        img = Image.open(file.stream).convert('RGB')
        images.append(img)

    if not images:
        return "No valid images uploaded", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

    # Save input
    filename = secure_filename(file.filename)
    input_path = os.path.join(upload_folder, filename)

    file.save(input_path)




    filename = f"imagepdf_{uuid.uuid4().hex}.pdf"
    output_path = os.path.join(app.config['SAVED_FOLDER'], filename)

    images[0].save(output_path, save_all=True, append_images=images[1:])
    return redirect(url_for('download_file', filename=filename))

#def compress_pdf(input_path: str, output_path: str, dpi: int = 72, quality: int = 50):
#    doc = fitz.open(input_path)
#    images = []
#    for page in doc:
#        pix = page.get_pixmap(dpi=dpi)
#        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#        img_byte = io.BytesIO()
#        img.save(img_byte, format="JPEG", quality=quality)
#        img_byte.seek(0)
#        images.append(Image.open(img_byte).convert("RGB"))

#    if not images:
#        raise ValueError("No pages found in the PDF to compress.")

#    images[0].save(output_path, save_all=True, append_images=images[1:])

@app.route('/download/<filename>')
def download_file(filename):
    path = os.path.join(app.config['SAVED_FOLDER'], filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    else:
        abort(404)


@app.route('/ocr_image_to_docx', methods=['POST'])
def ocr_image_to_docx():
    file = request.files['image']
    if file:
####
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 100 * 1024 * 1024:
            return "File size should not exceed 100 MB", 400

        upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

        # Save input
        filename = secure_filename(file.filename)
        input_path = os.path.join(upload_folder, filename)

        file.save(input_path)
####

        filename = f"ocr_{uuid.uuid4().hex}.docx"
        output_path = os.path.join(app.config['SAVED_FOLDER'], filename)

        img = Image.open(file.stream).convert('RGB')
        open_cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        text = pytesseract.image_to_string(open_cv_image)

        doc = Document()
        doc.add_paragraph(text)
        doc.save(output_path)

        return redirect(url_for('download_file', filename=filename))

#@app.route('/excel_to_pdf', methods=['POST'])
#def excel_to_pdf():
#    file = request.files['excel']
#    if file and (file.filename.endswith('.xls') or file.filename.endswith('.xlsx')):
#        file.seek(0, os.SEEK_END)
#        size = file.tell()
#        file.seek(0)
#        if size > 100 * 1024 * 1024:
#            return "File size should not exceed 100 MB", 400

#        upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])
#        filename = secure_filename(file.filename)
#        input_path = os.path.join(upload_folder, filename)
#        file.save(input_path)

#        output_filename = f"excelpdf_{uuid.uuid4().hex}.pdf"
#        output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)

#        try:
#            pythoncom.CoInitialize()
#            excel = comtypes.client.CreateObject('Excel.Application')
#            excel.Visible = False
#            wb = excel.Workbooks.Open(input_path)
#            wb.ExportAsFixedFormat(0, output_path)
#            wb.Close(False)
#            excel.Quit()
#        except Exception as e:
#            return f"Error converting Excel to PDF: {e}", 500

#        return redirect(url_for('download_file', filename=output_filename))
#    else:
#        return "Please upload a valid Excel file", 400







@app.route('/remove_pages', methods=['POST'])
def remove_pages():
    file = request.files['pdf']
    pages = request.form.get('pages')
    if not file or not pages:
        return "Please upload a PDF and specify pages to remove.", 400
#####
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 100 * 1024 * 1024:
        return "File size should not exceed 100 MB", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

    # Save input
    filename = secure_filename(file.filename)
    input_path = os.path.join(upload_folder, filename)

    file.save(input_path)




    try:
        pages_to_remove = set(int(p.strip()) - 1 for p in pages.split(',') if p.strip().isdigit())
        reader = PdfReader(file)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages):
            if i not in pages_to_remove:
                writer.add_page(page)

        output_filename = f"removed_pages_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)
        with open(output_path, 'wb') as f:
            writer.write(f)

        return redirect(url_for('download_file', filename=output_filename))

    except Exception as e:
        return f"Error removing pages: {e}", 500




@app.route('/rotate_pdf', methods=['POST'])
def rotate_pdf():
    file = request.files['pdf']
    rotations = request.form.get('rotations')
    if not file or not rotations:
        return "Please upload a PDF and specify rotations (e.g. 1:90,2:180).", 400

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 100 * 1024 * 1024:
        return "File size should not exceed 100 MB", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])

    # Save input
    filename = secure_filename(file.filename)
    input_path = os.path.join(upload_folder, filename)

    file.save(input_path)


    try:
        rotation_map = {}
        for entry in rotations.split(','):
            if ':' in entry:
                page, angle = entry.split(':')
                rotation_map[int(page.strip()) - 1] = int(angle.strip())

        reader = PdfReader(file)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages):
            if i in rotation_map:
                page.rotate(rotation_map[i])
            writer.add_page(page)

        output_filename = f"rotated_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(app.config['SAVED_FOLDER'], output_filename)
        with open(output_path, 'wb') as f:
            writer.write(f)

        return redirect(url_for('download_file', filename=output_filename))

    except Exception as e:
        return f"Error rotating PDF pages: {e}", 500


@app.route("/smart_split_merge")
def smart_split_merge_index():
    return SMART_SPLIT_TEMPLATE # Replace with actual HTML/template rendering

@app.route('/smart_split_merge/upload', methods=['POST'])
def smart_split_merge_upload():
    file = request.files['pdf']
    if not file:
        return jsonify({"error": "No file uploaded"}), 400

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 100 * 1024 * 1024:
        return "File size should not exceed 100 MB", 400

    upload_folder = create_timestamped_upload_folder(app.config['UPLOAD_FOLDER'])
    filename = secure_filename(file.filename)
    input_path = os.path.join(upload_folder, filename)
    file.save(input_path)

    reader = PdfReader(input_path)
    return jsonify({"filename": f"{os.path.basename(upload_folder)}/{filename}", "total_pages": len(reader.pages)})

@app.route('/smart_split_merge/generate', methods=['POST'])
def smart_split_merge_generate():
    data = request.json
    filename = data["filename"]
    selected_pages = data["pages"]
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    reader = PdfReader(input_path)
    generated_files = []
    for i in range(1, len(selected_pages) + 1):
        writer = PdfWriter()
        for p in selected_pages[:i]:
            writer.add_page(reader.pages[p - 1])
        out_name = f"merged_{'_'.join(map(str, selected_pages[:i]))}.pdf"
        out_path = os.path.join(app.config['SAVED_FOLDER'], out_name)
        with open(out_path, "wb") as f:
            writer.write(f)
        generated_files.append(out_name)
    return jsonify({"generated": generated_files})

@app.route('/smart_split_merge/uploads/<path:filename>')
def smart_split_merge_uploaded(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/smart_split_merge/generated/<path:filename>')
def smart_split_merge_generated(filename):
    return send_from_directory(app.config['SAVED_FOLDER'], filename)




@app.route('/download_saved/<filename>')
def download_saved_file(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(path):
        return send_file(path, as_attachment=True)
    else:
        abort(404)

@app.route('/saved_files')
def saved_files():
    files = os.listdir(app.config['SAVED_FOLDER'])
    links = [
        f'<li>{f} <a href="/download/{f}">Download</a> | <a href="/delete/{f}" style="color:red">Delete</a></li>'
        for f in files
    ]
    return '<h2>Saved Files</h2><ul>' + ''.join(links) + '</ul>'

@app.route('/delete/<filename>')
def delete_file(filename):
    path = os.path.join(app.config['SAVED_FOLDER'], filename)
    if os.path.exists(path):
        os.remove(path)
    return redirect(url_for('saved_files'))




#@app.route('/upload_folders')
#def list_upload_folders():
#   folders = os.listdir(app.config['UPLOAD_FOLDER'])
#   content = '<h2>Uploaded Folders</h2><ul>'
#   for folder in sorted(folders, reverse=True):
#       folder_path = os.path.join(app.config['UPLOAD_FOLDER'], folder)
#       if os.path.isdir(folder_path):
#           files = os.listdir(folder_path)
#           content += f'<li><strong>{folder}</strong><ul>'
#           for f in files:
#               file_path = f"/uploads/{folder}/{f}"
#               content += f'<li>{f} - <a href="{file_path}" target="_blank">View</a></li>'
#           content += '</ul></li>'
#   content += '</ul>'
#   return content

#@app.route('/uploads/<folder>/<filename>')
#def serve_uploaded_file(folder, filename):
#    folder_path = os.path.join(app.config['UPLOAD_FOLDER'], folder)
#    return send_file(os.path.join(folder_path, filename))


from flask import render_template_string

@app.route('/uploaded_files')
def uploaded_files():
    file_list = []
    base_path = app.config['UPLOAD_FOLDER']

    for root, dirs, files in os.walk(base_path):
        for name in files:
            # Get relative path for URL routing
            full_path = os.path.join(root, name)
            relative_path = os.path.relpath(full_path, base_path).replace("\\", "/")
            file_list.append(relative_path)

    return render_template_string("""
    <!doctype html>
    <title>Uploaded Files</title>
    <h1>Uploaded Files</h1>
    <ul>
    {% for file in files %}
      <li><a href="{{ url_for('serve_uploaded_file', filename=file) }}">{{ file }}</a></li>
    {% endfor %}
    </ul>
    """, files=file_list)
@app.route('/uploads/<path:filename>')
def serve_uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


#@app.route('/uploads/<folder>/<filename>')
#def serve_uploaded_file(folder, filename):
 #   return send_from_directory(os.path.join(app.config['UPLOAD_FOLDER'], folder), filename)





if __name__ == '__main__':
    app.run(debug=True)


